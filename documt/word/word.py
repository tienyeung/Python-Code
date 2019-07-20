import docx

#read the docx
def getText(filename):
    doc=docx.Document(filename)
    fulltext=[]
    for para in doc.paragraphs:
        fulltext.append(para.text)
    return '\n'.join(fulltext)

#write to docx
doc=docx.Document()
doc.add_paragraph('hello,world')
add title ,0 means 1st title
doc.add_heading('str',0)
换行：add_break 换页:add_break(docx.text.WD_BREAK.PAGE)
doc.save('hello_world.docx')


